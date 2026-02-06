from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('Best Buy Finder - Micro Project Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. Abstract
    doc.add_heading('1. Abstract', level=1)
    
    doc.add_heading('a. Brief Problem Statement', level=2)
    doc.add_paragraph(
        "Online shoppers often face difficulty in finding the lowest price for a product because prices vary across "
        "multiple e-commerce websites. Manually checking each website (Amazon, Best Buy, Walmart, etc.) is time-consuming and inefficient."
    )
    
    doc.add_heading('b. Objective of the Micro Project', level=2)
    doc.add_paragraph(
        "The objective of this project is to develop a web-based application that compares product prices across different "
        "e-commerce platforms in real-time and displays the lowest available price to users."
    )
    
    doc.add_heading('c. Tools / Technologies Used', level=2)
    doc.add_paragraph(
        "Python, Flask framework, MySQL database, SerpAPI (Google Shopping), HTML, CSS, Bootstrap, and concurrent processing."
    )
    
    doc.add_heading('d. Expected Outcome', level=2)
    doc.add_paragraph(
        "The system provides users with a simple interface to search products, view real-time prices from multiple vendors, "
        "and identify the best available price, saving time and money."
    )
    
    # 2. Introduction
    doc.add_heading('2. Introduction', level=1)
    
    doc.add_heading('a. Background of the Topic', level=2)
    doc.add_paragraph(
        "With the rapid growth of e-commerce, multiple online platforms sell the same products at different prices. "
        "Consumers need an efficient way to compare prices without visiting multiple websites individually."
    )
    
    doc.add_heading('b. Importance of the Problem', level=2)
    doc.add_paragraph(
        "Price comparison helps users make cost-effective purchasing decisions and promotes transparency in online shopping. "
        "It empowers consumers to get the best value for their money."
    )
    
    doc.add_heading('c. Real-Life Relevance', level=2)
    doc.add_paragraph(
        "Most consumers compare prices before buying products such as electronics, fashion items, and household goods. "
        "This system automates that comparison by aggregating data from major retailers like Amazon, Walmart, Best Buy, and more."
    )
    
    doc.add_heading('d. Scope of the Micro Project', level=2)
    doc.add_paragraph(
        "The project focuses on comparing prices from selected e-commerce websites using real-time API data. "
        "It includes features like user authentication, shopping cart, price history tracking, and price alerts."
    )
    
    # 3. Problem Definition
    doc.add_heading('3. Problem Definition', level=1)
    
    doc.add_heading('a. Clear Problem Statement', level=2)
    doc.add_paragraph(
        "To design and implement a web application that compares product prices from different e-commerce websites "
        "and displays the lowest price to users dynamically."
    )
    
    doc.add_heading('b. Existing System', level=2)
    doc.add_paragraph(
        "Users manually search for products on multiple websites like Amazon, Flipkart, etc., and compare prices themselves. "
        "Some existing comparison sites may have outdated data."
    )
    
    doc.add_heading('c. Limitations of Existing System', level=2)
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run("Time-consuming process to check sites one by one.")
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run("Possibility of missing better deals.")
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run("No centralized cart or alert system.")
    
    doc.add_heading('d. Proposed Solution (Overview)', level=2)
    doc.add_paragraph(
        "A centralized web application that fetches product data via API, compares prices in real-time, "
        "and displays the best buy option with features like 'Best Deal' highlighting."
    )
    
    # 4. Objectives
    doc.add_heading('4. Objectives of the Micro Project', level=1)
    objectives = [
        "To understand the basics of web development using Python Flask.",
        "To implement a price comparison system using real-time API integration (SerpAPI).",
        "To manage user data and orders using a MySQL database.",
        "To optimize performance using parallel processing for API requests.",
        "To improve user experience using a responsive UI with Bootstrap."
    ]
    for obj in objectives:
        p = doc.add_paragraph(obj)
        p.style = 'List Bullet'

    # 5. Methodology
    doc.add_heading('5. Methodology / Working', level=1)
    
    doc.add_heading('a. Data Collection', level=2)
    doc.add_paragraph(
        "Product data (name, price, image, source) is fetched in real-time using SerpAPI, which aggregates results from Google Shopping. "
        "Data is normalized to a common format and cached for 24 hours to improve performance."
    )
    
    doc.add_heading('b. Algorithm / Process Flow', level=2)
    steps = [
        "User enters product name.",
        "System checks local cache for recent results.",
        "If not cached, system launches parallel threads to fetch data from multiple sources (Amazon, Best Buy, etc.) simultaneously.",
        "Data is collected, normalized, and sorted by price.",
        "Lowest price product is identified and highlighted.",
        "Result is displayed to the user."
    ]
    for step in steps:
        p = doc.add_paragraph(step, style='List Number')
        
    doc.add_heading('c. System Architecture', level=2)
    components = [
        "Frontend: HTML, CSS, Bootstrap (Responsive UI)",
        "Backend: Flask (Python Web Server)",
        "Database: MySQL (User storage, Orders, Alerts)",
        "External API: SerpAPI (Google Shopping Data)"
    ]
    for comp in components:
        doc.add_paragraph(comp, style='List Bullet')
        
    doc.add_heading('d. Flowchart / Block Diagram', level=2)
    doc.add_paragraph("User → Web Interface → Flask Server (Parallel Fetching) → SerpAPI/MySQL → Price Comparison Logic → Result Display")
    
    # 6. Tools and Technologies
    doc.add_heading('6. Tools and Technologies Used', level=1)
    
    doc.add_heading('a. Hardware', level=2)
    doc.add_paragraph("Computer / Laptop, Internet Connection")
    
    doc.add_heading('b. Software', level=2)
    softwares = [
        "Python (Programming Language)",
        "Flask Framework (Backend)",
        "MySQL (Database)",
        "Visual Studio Code / PyCharm (IDE)",
        "Web Browser"
    ]
    for s in softwares:
        doc.add_paragraph(s, style='List Bullet')
        
    doc.add_heading('c. Libraries / Packages', level=2)
    libs = [
        "Flask (Web Server)",
        "mysql-connector-python (Database Driver)",
        "requests (API Calls)",
        "concurrent.futures (Parallel Processing)",
        "Bootstrap (UI Design)"
    ]
    for l in libs:
        doc.add_paragraph(l, style='List Bullet')

    # 7. Implementation
    doc.add_heading('7. Implementation', level=1)
    
    doc.add_heading('a. Input Description', level=2)
    doc.add_paragraph("User inputs product name or category through a search form (e.g., 'iPhone 15', 'Running Shoes').")
    
    doc.add_heading('b. Processing Steps', level=2)
    proc_steps = [
        "Flask receives input query.",
        "ThreadPoolExecutor initiates simultaneous requests to various store endpoints.",
        "Responses are parsed to extract price, title, and image.",
        "Prices are converted to INR (if needed) and sorted ascending.",
        "Best deal is calculated."
    ]
    for ps in proc_steps:
        doc.add_paragraph(ps, style='List Bullet')
        
    doc.add_heading('c. Output Description', level=2)
    doc.add_paragraph(
        "The application displays a grid of product cards from different retailers. "
        "A special 'Lowest Price' card highlights the best deal found."
    )
    
    doc.add_heading('d. Screenshots / Sample Data', level=2)
    doc.add_paragraph("(Please insert screenshots of the Homepage, Search Results, and Lowest Price Card here)")
    
    # 8. Result and Analysis
    doc.add_heading('8. Result and Analysis', level=1)
    
    doc.add_heading('a. Output Obtained', level=2)
    doc.add_paragraph("The system successfully compares product prices from Amazon, Best Buy, Walmart, and others, displaying the cheapest option at the top.")
    
    doc.add_heading('b. Observations', level=2)
    obs = [
        "Faster product comparison due to parallel processing.",
        "Clean and user-friendly interface.",
        "Accurate real-time pricing from live markets."
    ]
    for o in obs:
        doc.add_paragraph(o, style='List Bullet')
        
    doc.add_heading('c. Limitations', level=2)
    doc.add_paragraph("Dependent on third-party API availability and rate limits (SerpAPI free tier limitations).")
    
    # 9. Applications
    doc.add_heading('9. Applications', level=1)
    apps = [
        "Industry Applications: Used by e-commerce aggregators and online marketplaces.",
        "Educational Use: Helps students understand full-stack web development and API integration.",
        "Business Use: Useful for customers and retailers to monitor pricing strategies.",
        "Future Scope: Mobile App development, AI-based price prediction, and barcode scanning."
    ]
    for app in apps:
        doc.add_paragraph(app, style='List Bullet')
        
    # 10. SDG
    doc.add_heading('10. Relevance to Sustainable Development Goals', level=1)
    doc.add_paragraph(
        "This project supports SDG 12: Responsible Consumption and Production by helping consumers make informed "
        "and economical purchasing decisions, reducing unnecessary spending."
    )
    
    # 11. Conclusion
    doc.add_heading('11. Conclusion', level=1)
    
    doc.add_heading('A. Summary of Work Done', level=2)
    doc.add_paragraph(
        "A price comparison e-commerce web application was successfully developed using Python, Flask, MySQL, and Bootstrap. "
        "It integrates real-time data fetching and user management."
    )
    
    doc.add_heading('B. Learning Outcomes', level=2)
    outcomes = [
        "Understanding of Flask framework and MVC architecture.",
        "Database integration using MySQL.",
        "Consuming REST APIs and handling JSON data.",
        "Parallel programming for performance optimization.",
        "UI design using Bootstrap."
    ]
    for out in outcomes:
        doc.add_paragraph(out, style='List Bullet')
        
    doc.add_heading('C. Achievement of Objectives', level=2)
    doc.add_paragraph("All project objectives were successfully achieved, and the system performs efficiently as expected.")

    # Save
    doc.save('Project_Report.docx')
    print("Report generated successfully: Project_Report.docx")

if __name__ == "__main__":
    create_report()
